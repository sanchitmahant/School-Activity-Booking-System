"""
Professional DOCX Generator for CN7021 Final Report
Matches exact format from reference document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_professional_docx():
    """Main function to create professionally formatted DOCX"""
    print("🔄 Creating professional DOCX from markdown...")
    
    # Create new document
    doc = Document()
    
    # Set default font for Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # TITLE PAGE
    print("Creating title page...")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SCHOOL OF ARCHITECTURE, COMPUTING & ENGINEERING')
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CN7021 Advanced Software Engineering')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('School Activity Booking System')
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Student details
    p = doc.add_paragraph('Student Name: Sanchit Kaushal')
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
   
    p = doc.add_paragraph('Student ID: 2823183')
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    
    p = doc.add_paragraph('Tutor Name: [TUTOR NAME]')
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    
    p = doc.add_paragraph('Lab Number: [LAB NUMBER]')
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('Submission Date: 22nd December 2025')
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    
    # Page break
    doc.add_page_break()
    
    # TABLE OF CONTENTS
    print("Creating TOC placeholder...")
    p = doc.add_heading('Table of Contents', 1)
    p = doc.add_paragraph('[INSERT TABLE OF CONTENTS HERE]')
    p = doc.add_paragraph('In Word: References → Table of Contents → Automatic Table')
    p.runs[0].font.italic = True
    p.runs[0].font.size = Pt(10)
    
    doc.add_page_break()
    
    # Read markdown content
    print("Processing markdown content...")
    with open('FINAL_REPORT_DRAFT.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Process content line by line
    lines = content.split('\n')
    
    in_table = False
    table_lines = []
    
    for line in lines:
        line = line.rstrip()
        
        # Handle tables
        if '|' in line and line.strip() and not line.strip().startswith('['):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        else:
            if in_table and table_lines:
                # Create table
                create_table_from_markdown(doc, table_lines)
                in_table = False
                table_lines = []
        
        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            continue
        
        # Bullet points
        if line.startswith('*   ') or line.startswith('- '):
            text = line[4:] if line.startswith('*   ') else line[2:]
            # Remove markdown
            text = remove_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
            continue
        
        # Screenshot placeholders
        if '[INSERT' in line and 'HERE]' in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.bold = True
            try:
                from docx.shared import RGBColor
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            except:
                pass
            continue
        
        # Figure captions
        if line.startswith('Figure '):
            p = doc.add_paragraph(line)
            p.runs[0].font.italic = True
            p.runs[0].font.size = Pt(10)
            continue
        
        # Normal paragraph
        text = remove_markdown(line)
        if text.strip():
            p = doc.add_paragraph(text)
    
    # Handle remaining table
    if table_lines:
        create_table_from_markdown(doc, table_lines)
    
    # Save
    output_file = 'Sanchit_Kaushal_Final_Report_PROFESSIONAL.docx'
    doc.save(output_file)
    
    # Stats
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"✅ Successfully generated: {output_file}")
    print(f"📄 Paragraphs: {len(doc.paragraphs)}")
    print(f"📊 Tables: {len(doc.tables)}")
    print(f"📝 Word count: {words}")
    print(f"✅ Status: {'COMPLIANT' if 2700 <= words <= 3300 else 'NEEDS ADJUSTMENT'}")

def remove_markdown(text):
    """Remove markdown formatting"""
    # Bold italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Code
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

def create_table_from_markdown(doc, lines):
    """Create Word table from markdown table lines"""
    # Parse rows
    rows = []
    for line in lines:
        if '---' in line or not line.strip():
            continue
        cells = [cell.strip() for cell in line.split('|')]
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)
    
    if not rows or len(rows) < 1:
        return
    
    # Create table
    try:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        # Fill cells
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    cell.text = remove_markdown(cell_text)
                    
                    # Bold header
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
        
        doc.add_paragraph()  # Spacing
    except Exception as e:
        print(f"⚠️ Table creation error: {e}")
        # Fallback: add as paragraph
        for row in rows:
            doc.add_paragraph(' | '.join(row))

if __name__ == '__main__':
    create_professional_docx()
