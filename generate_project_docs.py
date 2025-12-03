import os
import ast
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = os.getcwd()
DOCS_DIR = os.path.join(PROJECT_ROOT, "documentation", "Folder and FIles documentation")
IGNORE_DIRS = {'.git', '.venv', '__pycache__', 'documentation', '.idea', '.vscode', 'instance'}
IGNORE_FILES = {'.gitignore', '.env', '.DS_Store', 'generate_project_docs.py', 'read_docx.py', 'extract_docx_content.py', 'view_csv.py', 'view_csv_detailed.py', 'analyze_distribution.py', 'redistribute_tasks_equal.py', 'check_equal_distribution.py', 'distribute_tasks.py'}

# Pre-defined explanations for key files
FILE_EXPLANATIONS = {
    "app.py": {
        "simple": "This is the main brain of the application. It starts the website, connects to the database (where data is stored), and decides what page to show when you click on links. It also defines what a 'User', 'Activity', or 'Booking' looks like.",
        "technical": "Entry point for the Flask application. Initializes the Flask app context, configures SQLAlchemy for ORM-based database interactions, and defines the database models (Schema). It also contains the route definitions (Controllers) that handle HTTP requests and render Jinja2 templates."
    },
    "config.py": {
        "simple": "This file stores the settings for the project, like passwords and secret keys. It's like the settings menu on your phone, but for the code.",
        "technical": "Configuration module containing class-based config objects (Config, DevelopmentConfig, ProductionConfig). It manages environment variables, database URIs, and security keys to ensure separation of configuration from code."
    },
    "populate_db.py": {
        "simple": "This script fills the database with fake data so you have something to look at when you first run the app. It creates sample users, activities, and bookings.",
        "technical": "Database seeding script. It initializes the application context and uses SQLAlchemy sessions to insert initial test data (fixtures) into the database, ensuring a consistent state for development and testing."
    },
    "requirements.txt": {
        "simple": "A list of all the extra tools (libraries) this project needs to run. It's like a shopping list for the computer.",
        "technical": "Dependency manifest file listing all Python packages and their specific versions required to run the application, used by pip for reproducible environment setup."
    },
    "base.html": {
        "simple": "The master template. It contains the parts of the website that stay the same on every page, like the navigation bar at the top and the footer at the bottom.",
        "technical": "Base Jinja2 template that implements the common HTML structure (DOCTYPE, head, scripts). It defines 'blocks' that child templates can override, enforcing DRY (Don't Repeat Yourself) principles in frontend code."
    },
    "index.html": {
        "simple": "The home page of the website. This is the first thing users see.",
        "technical": "Landing page template. Extends 'base.html' and displays the primary call-to-action and introductory content."
    },
    "login.html": {
        "simple": "The page where users enter their email and password to access their account.",
        "technical": "Authentication template containing the login form. It submits POST requests to the login route and displays flash messages for errors or success."
    },
    "style.css": {
        "simple": "Controls how the website looks—colors, fonts, spacing, and layout.",
        "technical": "Global Cascading Style Sheet (CSS). Defines the visual styling, responsive grid adjustments, and custom component designs (buttons, cards, forms) to ensure UI consistency."
    },
    "script.js": {
        "simple": "Handles the interactive parts of the website, like pop-ups or checking if you filled out a form correctly.",
        "technical": "Client-side JavaScript file. Implements DOM manipulation, event listeners for UI interactions, and potentially AJAX calls for asynchronous data updates."
    }
}

GENERIC_EXPLANATIONS = {
    ".py": {
        "simple": "A Python code file containing logic for the application.",
        "technical": "Python source module containing classes, functions, and business logic."
    },
    ".html": {
        "simple": "A web page template that defines the structure of the content.",
        "technical": "HTML template file, likely using Jinja2 syntax for dynamic content rendering."
    },
    ".css": {
        "simple": "A styling file that makes the web pages look good.",
        "technical": "CSS stylesheet for defining presentation and layout."
    },
    ".js": {
        "simple": "A script file that makes the web page interactive.",
        "technical": "JavaScript source file for client-side logic."
    }
}

def create_document(title):
    doc = Document()
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    runner = p.add_run(text)
    if bold:
        runner.bold = True
    if italic:
        runner.italic = True
    if color:
        runner.font.color.rgb = color

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    runner = p.add_run(code)
    runner.font.name = 'Courier New'
    runner.font.size = Pt(9)

def analyze_python_detailed(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
            lines = content.splitlines()
        
        tree = ast.parse(content)
        
        analysis = {
            'docstring': ast.get_docstring(tree),
            'imports': [],
            'classes': [],
            'functions': [],
            'constants': [],
            'full_content': content
        }
        
        for node in ast.iter_child_nodes(tree):
            if isinstance(node, (ast.Import, ast.ImportFrom)):
                if isinstance(node, ast.Import):
                    names = [n.name for n in node.names]
                    analysis['imports'].append(f"Import: {', '.join(names)}")
                else:
                    module = node.module or ''
                    names = [n.name for n in node.names]
                    analysis['imports'].append(f"From {module} import {', '.join(names)}")
            
            elif isinstance(node, ast.Assign):
                try:
                    targets = [t.id for t in node.targets if isinstance(t, ast.Name)]
                    if targets:
                        val = "value"
                        if isinstance(node.value, ast.Constant):
                            val = str(node.value.value)
                        analysis['constants'].append(f"{', '.join(targets)} = {val}")
                except:
                    pass

            elif isinstance(node, ast.ClassDef):
                class_info = {
                    'name': node.name,
                    'docstring': ast.get_docstring(node),
                    'methods': [],
                    'bases': [b.id for b in node.bases if isinstance(b, ast.Name)]
                }
                for item in node.body:
                    if isinstance(item, ast.FunctionDef):
                        args = [a.arg for a in item.args.args]
                        class_info['methods'].append({
                            'name': item.name,
                            'args': args,
                            'docstring': ast.get_docstring(item),
                            'code': "\n".join(lines[item.lineno-1:item.end_lineno])
                        })
                analysis['classes'].append(class_info)

            elif isinstance(node, ast.FunctionDef):
                args = [a.arg for a in node.args.args]
                analysis['functions'].append({
                    'name': node.name,
                    'args': args,
                    'docstring': ast.get_docstring(node),
                    'code': "\n".join(lines[node.lineno-1:node.end_lineno])
                })
                
        return analysis
    except Exception as e:
        return {'error': str(e), 'full_content': content if 'content' in locals() else ''}

def generate_file_doc(filepath, rel_path, output_dir):
    filename = os.path.basename(filepath)
    if filename in IGNORE_FILES:
        return

    doc_name = f"{filename}_Documentation.docx"
    doc_path = os.path.join(output_dir, doc_name)
    
    doc = create_document(f"Documentation: {filename}")
    
    # 1. File Overview
    add_heading(doc, "1. File Overview", level=1)
    add_paragraph(doc, f"File Name: {filename}", bold=True)
    add_paragraph(doc, f"Location: {os.path.dirname(rel_path)}")
    add_paragraph(doc, f"Type: {os.path.splitext(filename)[1]}")

    # 2. Explanations (Simple & Technical)
    add_heading(doc, "2. Explanation", level=1)
    
    # Get explanation
    explanation = FILE_EXPLANATIONS.get(filename)
    if not explanation:
        ext = os.path.splitext(filename)[1].lower()
        explanation = GENERIC_EXPLANATIONS.get(ext, {
            "simple": "A project file.",
            "technical": "Source file."
        })

    # Simple Explanation
    add_heading(doc, "Simple Language (For Non-Tech)", level=2)
    add_paragraph(doc, explanation["simple"], italic=True)
    
    # Technical Explanation
    add_heading(doc, "Technical Language (For Developers)", level=2)
    add_paragraph(doc, explanation["technical"])

    ext = os.path.splitext(filename)[1].lower()
    
    if ext == '.py':
        analysis = analyze_python_detailed(filepath)
        
        if 'error' in analysis:
            add_paragraph(doc, f"Error analyzing file: {analysis['error']}")
        else:
            # 3. Code Structure
            add_heading(doc, "3. Code Structure & Logic", level=1)
            
            # Imports
            if analysis.get('imports'):
                add_heading(doc, "Dependencies", level=2)
                for imp in analysis['imports']:
                    add_paragraph(doc, f"• {imp}")

            # Classes
            if analysis.get('classes'):
                add_heading(doc, "Classes", level=2)
                for cls in analysis['classes']:
                    bases = f"({', '.join(cls['bases'])})" if cls['bases'] else ""
                    add_paragraph(doc, f"Class: {cls['name']}{bases}", bold=True, color=RGBColor(0, 0, 139))
                    if cls['docstring']:
                        add_paragraph(doc, f"Docstring: {cls['docstring']}")
                    
                    if cls['methods']:
                        for method in cls['methods']:
                            args = ", ".join(method['args'])
                            add_paragraph(doc, f"  Method: {method['name']}({args})", bold=True)
                            if method['docstring']:
                                add_paragraph(doc, f"  - {method['docstring']}")

            # Functions
            if analysis.get('functions'):
                add_heading(doc, "Functions", level=2)
                for func in analysis['functions']:
                    args = ", ".join(func['args'])
                    add_paragraph(doc, f"Function: {func['name']}({args})", bold=True, color=RGBColor(0, 100, 0))
                    if func['docstring']:
                        add_paragraph(doc, f"Purpose: {func['docstring']}")
                    
                    # Add code snippet for logic
                    add_paragraph(doc, "Logic Preview:")
                    add_code_block(doc, func['code'][:500] + "..." if len(func['code']) > 500 else func['code'])

            # Full Code
            add_heading(doc, "4. Full Source Code", level=1)
            add_code_block(doc, analysis['full_content'])

    else:
        # Non-Python files
        add_heading(doc, "3. File Content", level=1)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            add_paragraph(doc, f"Size: {len(content)} bytes")
            add_paragraph(doc, "Content Preview:")
            add_code_block(doc, content)
        except Exception as e:
            add_paragraph(doc, f"Could not read file content: {e}")

    doc.save(doc_path)
    print(f"Generated: {rel_path}")

def generate_folder_doc(folder_path, rel_path, output_dir):
    folder_name = os.path.basename(folder_path)
    if folder_name in IGNORE_DIRS:
        return

    doc_name = f"Folder_{folder_name}_Overview.docx"
    doc_path = os.path.join(output_dir, doc_name)
    
    doc = create_document(f"Folder: {folder_name}")
    
    add_heading(doc, "1. Folder Overview", level=1)
    add_paragraph(doc, f"Name: {folder_name}")
    add_paragraph(doc, f"Path: {rel_path}")
    
    add_heading(doc, "2. Explanation", level=1)
    add_paragraph(doc, "Simple: A container for organizing related files.", italic=True)
    add_paragraph(doc, "Technical: Directory structure for module organization and separation of concerns.")
    
    try:
        items = os.listdir(folder_path)
        files = [f for f in items if os.path.isfile(os.path.join(folder_path, f)) and f not in IGNORE_FILES]
        subdirs = [d for d in items if os.path.isdir(os.path.join(folder_path, d)) and d not in IGNORE_DIRS]
        
        add_heading(doc, "3. Contents", level=1)
        if subdirs:
            add_heading(doc, "Subdirectories", level=2)
            for d in subdirs:
                add_paragraph(doc, f"📂 {d}/")
        
        if files:
            add_heading(doc, "Files", level=2)
            for f in files:
                add_paragraph(doc, f"📄 {f}")
                
    except Exception as e:
        add_paragraph(doc, f"Error: {e}")
        
    doc.save(doc_path)

def main():
    if os.path.exists(DOCS_DIR):
        shutil.rmtree(DOCS_DIR)
    os.makedirs(DOCS_DIR)
    
    print(f"Scanning: {PROJECT_ROOT}")
    generate_folder_doc(PROJECT_ROOT, ".", DOCS_DIR)
    
    for root, dirs, files in os.walk(PROJECT_ROOT):
        dirs[:] = [d for d in dirs if d not in IGNORE_DIRS]
        rel_root = os.path.relpath(root, PROJECT_ROOT)
        
        if root == PROJECT_ROOT:
            current_output_dir = DOCS_DIR
        else:
            current_output_dir = os.path.join(DOCS_DIR, rel_root)
            if not os.path.exists(current_output_dir):
                os.makedirs(current_output_dir)
            generate_folder_doc(root, rel_root, current_output_dir)
            
        for file in files:
            if file not in IGNORE_FILES:
                filepath = os.path.join(root, file)
                rel_path = os.path.join(rel_root, file)
                generate_file_doc(filepath, rel_path, current_output_dir)

if __name__ == "__main__":
    main()
