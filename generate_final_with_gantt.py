"""
Enhanced DOCX Generator with Gantt Chart Embedded
Final submission version
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
import os

def remove_markdown(text):
    """Remove all markdown formatting"""
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text

def create_professional_docx_with_gantt():
    """Main function with Gantt chart"""
    print("🔄 Generating FINAL submission DOCX with Gantt chart...")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # TITLE PAGE
    print("Creating title page...")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SCHOOL OF ARCHITECTURE, COMPUTING & ENGINEERING')
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department of Computer Science and Digital Technologies – CDT')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CN7021 – Advanced Software Engineering')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('School Activity Booking System')
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Group: 3.B')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Student names
    students = [
        'Sanchit Kaushal (2823183)',
        'Mohd Sharjeel (2823311)',
        'Chichebendu Umeh (2823112)',
        'Shiva Kasula (2822121)'
    ]
    for student in students:
        p = doc.add_paragraph(student)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'Calibri'
        p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph('Tutor: [TUTOR NAME]')
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    
    p = doc.add_paragraph('Module Leader: Dr Hisham AbouGrad')
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('December 2025')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    
    # Page break
    doc.add_page_break()
    
    # TABLE OF CONTENTS
    print("Creating TOC placeholder...")
    p = doc.add_heading('Table of Contents', 1)
    p = doc.add_paragraph('[INSERT TABLE OF CONTENTS HERE]')
    p = doc.add_paragraph('In Word: References → Table of Contents → Automatic Table')
    p.runs[0].font.italic = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # Read markdown content
    print("Processing markdown content...")
    with open('FINAL_REPORT_DRAFT.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Skip title page content in markdown
    content_start = False
    for i, line in enumerate(lines):
        if line.strip() == '1. Introduction':
            content_start = True
            break
    
    if content_start:
        lines = lines[i:]
    
    # Process content and INSERT GANTT CHART
    in_table = False
    table_lines = []
    gantt_inserted = False
    
    for line in lines:
        line = line.rstrip()
        
        # INSERT GANTT CHART after Section 8.3 heading
        if not gantt_inserted and '8.3 Project Management Board' in line:
            # Add heading
            p = doc.add_heading(line, level=2)
            
            # Add Gantt chart image if it exists
            gantt_path = 'gantt_chart_project.webp'
            if os.path.exists(gantt_path):
                print(f"📊 Inserting Gantt chart: {gantt_path}")
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(gantt_path, width=Inches(6.5))
                
                # Add caption
                p = doc.add_paragraph('Figure 8: Project Gantt Chart - Sprint Timeline (Oct-Dec 2024)')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.italic = True
                p.runs[0].font.size = Pt(10)
                p = doc.add_paragraph('(Source: Project Management Documentation)')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.italic = True
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
                doc.add_paragraph()
                gantt_inserted = True
            continue
        
        # Handle tables
        if '|' in line and line.strip() and not line.strip().startswith('['):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        else:
            if in_table and table_lines:
                create_table(doc, table_lines)
                in_table = False
                table_lines = []
        
        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Numbered sections
        if re.match(r'^\d+\.\s+[A-Z]', line):
            p = doc.add_heading(line, level=1)
            continue
        elif re.match(r'^\d+\.\d+\s+[A-Z]', line):
            p = doc.add_heading(line, level=2)
            continue
        elif re.match(r'^\d+\.\d+\.\d+\s+[A-Z]', line):
            p = doc.add_heading(line, level=3)
            continue
        
        # Bullet points
        if line.startswith('*   ') or line.startswith('- '):
            text = line[4:] if line.startswith('*   ') else line[2:]
            text = remove_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
            continue
        
        # INSERT placeholders
        if '[INSERT' in line and 'HERE]' in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            continue
        
        # Figure captions
        if line.startswith('Figure '):
            p = doc.add_paragraph(line)
            p.runs[0].font.italic = True
            p.runs[0].font.size = Pt(10)
            continue
        
        # Normal paragraph
        text = remove_markdown(line)
        if text.strip():
            p = doc.add_paragraph(text)
    
    # Handle remaining table
    if table_lines:
        create_table(doc, table_lines)
    
    # Save
    output_file = 'Sanchit_Kaushal_Final_Report_SUBMISSION.docx'
    doc.save(output_file)
    
    # Stats
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"✅ Successfully generated: {output_file}")
    print(f"📄 Paragraphs: {len(doc.paragraphs)}")
    print(f"📊 Tables: {len(doc.tables)}")
    print(f"📝 Word count: {words}")
    print(f"📊 Gantt chart: {'EMBEDDED ✅' if gantt_inserted else 'Not found'}")
    status = 'COMPLIANT ✅' if 2700 <= words <= 3400 else f'CHECK MANUALLY'
    print(f"✅ Status: {status}")

def create_table(doc, lines):
    """Create Word table from markdown table lines"""
    rows = []
    for line in lines:
        if '---' in line or not line.strip():
            continue
        cells = [cell.strip() for cell in line.split('|')]
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)
    
    if not rows or len(rows) < 1:
        return
    
    try:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    cell.text = remove_markdown(cell_text)
                    
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
        
        doc.add_paragraph()
    except Exception as e:
        print(f"⚠️ Table creation error: {e}")

if __name__ == '__main__':
    create_professional_docx_with_gantt()
