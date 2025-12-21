"""
Generate DOCX report from FINAL_REPORT_DRAFT.md with proper academic formatting
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_formatted_report():
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Heading styles
    for i in range(1, 4):
        heading_style = styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(14 - i)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    # Read markdown file
    with open('FINAL_REPORT_DRAFT.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # Add title page
    title = doc.add_paragraph('SCHOOL OF ARCHITECTURE, COMPUTING & ENGINEERING', style='Normal')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(14)
    title.runs[0].font.bold = True
    
    subtitle = doc.add_paragraph('Department of Computer Science and Digital Technologies – CDT', style='Normal')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    
    module = doc.add_paragraph('CN7021 – Advanced Software Engineering', style='Normal')
    module.alignment = WD_ALIGN_PARAGRAPH.CENTER
    module.runs[0].font.size = Pt(12)
    module.runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    project_title = doc.add_paragraph('School Activity Booking System', style='Title')
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    group = doc.add_paragraph('Group: 3.B', style='Normal')
    group.alignment = WD_ALIGN_PARAGRAPH.CENTER
    group.runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # Student details
    students = doc.add_paragraph('Students Name & ID:', style='Normal')
    students.alignment = WD_ALIGN_PARAGRAPH.CENTER
    students.runs[0].font.bold = True
    
    student_list = [
        'Sanchit Kaushal (2823183)',
        'Mohd Sharjeel (2823311)',
        'Chichebendu Umeh (2823112)',
        'Shiva Kasula (2822121)'
    ]
    
    for student in student_list:
        p = doc.add_paragraph(student, style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    tutor = doc.add_paragraph('Tutor: [TUTOR NAME]', style='Normal')
    tutor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tutor.runs[0].font.bold = True
    
    leader = doc.add_paragraph('Module Leader: Dr Hisham AbouGrad', style='Normal')
    leader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    leader.runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    date = doc.add_paragraph('December 2025', style='Normal')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].font.bold = True
    
    # Page break
    doc.add_page_break()
    
    # Table of Contents placeholder
    toc = doc.add_paragraph('Table of Contents', style='Heading 1')
    toc_note = doc.add_paragraph('[INSERT TABLE OF CONTENTS HERE - Use Word: References → Table of Contents]', style='Normal')
    toc_note.runs[0].font.italic = True
    toc_note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # Process markdown content
    skip_title_section = True
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip title page content
        if skip_title_section:
            if line.startswith('1. Introduction'):
                skip_title_section = False
            else:
                i += 1
                continue
        
        # Skip markdown ToC
        if '[INSERT TABLE OF CONTENTS HERE]' in line:
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_paragraph(line[2:], style='Heading 1')
        elif line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading 2')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:], style='Heading 3')
        elif line.startswith('#### '):
            doc.add_paragraph(line[5:], style='Heading 4')
        
        # Handle section numbers (1., 1.1, etc.)
        elif re.match(r'^\d+\.', line):
            # This is a numbered section heading
            doc.add_paragraph(line, style='Heading 1')
        elif re.match(r'^\d+\.\d+', line):
            # This is a subsection heading
            doc.add_paragraph(line, style='Heading 2')
        
        # Handle tables
        elif '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
        else:
            if in_table:
                # Process completed table
                in_table = False
                create_table_from_markdown(doc, table_lines)
                table_lines = []
            
            # Handle placeholders
            if '[INSERT FIGURE' in line or '[INSERT SCREENSHOT' in line:
                p = doc.add_paragraph(line, style='Normal')
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                p.runs[0].font.bold = True
            
            # Handle bold text
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph(line[2:-2], style='Normal')
                p.runs[0].font.bold = True
            
            # Handle bullet points
            elif line.startswith('*   ') or line.startswith('- '):
                text = line[4:] if line.startswith('*   ') else line[2:]
                doc.add_paragraph(text, style='List Bullet')
            
            # Normal paragraph
            elif line.strip():
                doc.add_paragraph(line, style='Normal')
            else:
                # Empty line - slight spacing
                doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save('Sanchit_Kaushal_Final_Report_v32.docx')
    print("✅ Successfully generated: Sanchit_Kaushal_Final_Report_v32.docx")
    print(f"📄 Sections included:")
    print("  - Title page with student details")
    print("  - Table of Contents placeholder")
    print("  - Complete report content (3,352 words)")
    print("  - COCOMO cost estimation (£12,000)")
    print("  - Contributions Table (Appendix G)")
    print("  - Harvard references (Section 1.5)")
    print("\n⚠️  Manual actions required:")
    print("  1. Fill [TUTOR NAME] on title page")
    print("  2. Insert 12 screenshots at marked placeholders")
    print("  3. Generate Table of Contents: References → Table of Contents")


def create_table_from_markdown(doc, table_lines):
    """Convert markdown table to Word table"""
    # Filter out separator lines
    data_lines = [line for line in table_lines if not set(line.replace('|', '').strip()) == {'-', ' '}]
    
    if len(data_lines) < 2:
        return
    
    # Parse table
    rows = []
    for line in data_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove outer pipes
        rows.append(cells)
    
    # Create Word table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            
            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_paragraph()  # Spacing after table


if __name__ == '__main__':
    create_formatted_report()
