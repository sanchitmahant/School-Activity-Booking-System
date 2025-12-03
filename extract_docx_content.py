from docx import Document
import sys

def extract_content(filepath):
    """Extract all text content from a docx file"""
    doc = Document(filepath)
    content = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                content.append(' | '.join(row_data))
    
    return '\n'.join(content)

if __name__ == "__main__":
    announcement_path = r"documentation\Announcement.docx"
    brief_path = r"documentation\CN7021 ASWE 2025-26 Coursework Brief.docx"
    
    # Extract announcement content
    with open("announcement_content.txt", "w", encoding="utf-8") as f:
        f.write("=" * 80 + "\n")
        f.write("ANNOUNCEMENT CONTENT\n")
        f.write("=" * 80 + "\n")
        f.write(extract_content(announcement_path))
    
    # Extract coursework brief content
    with open("coursework_brief_content.txt", "w", encoding="utf-8") as f:
        f.write("=" * 80 + "\n")
        f.write("COURSEWORK BRIEF CONTENT\n")
        f.write("=" * 80 + "\n")
        f.write(extract_content(brief_path))
    
    print("Content extracted successfully!")
    print("- announcement_content.txt")
    print("- coursework_brief_content.txt")
