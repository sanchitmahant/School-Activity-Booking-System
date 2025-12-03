from docx import Document
import sys

def read_docx(filepath):
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also read tables if present
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(' | '.join(row_text))
    
    return '\n'.join(full_text)

if __name__ == "__main__":
    if len(sys.argv) > 1:
        filepath = sys.argv[1]
        content = read_docx(filepath)
        print(content)
    else:
        print("Please provide a docx file path")
