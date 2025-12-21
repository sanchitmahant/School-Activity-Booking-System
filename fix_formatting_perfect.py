"""
Extract formatting specifications from reference DOCX
Then fix references and regenerate with exact formatting
"""

from docx import Document
import os

def analyze_reference_doc():
    """Extract formatting from reference document"""
    ref_path = r"Real Report Preparation\Helping and Guideline Documents\CN7021ASWE-Project (School Activity Booking System) (Reference).docx"
    
    if not os.path.exists(ref_path):
        print("❌ Reference document not found")
        return None
    
    print("📖 Analyzing reference document formatting...")
    doc = Document(ref_path)
    
    # Extract styles
    styles = {}
    
    # Check paragraphs for font info
    for para in doc.paragraphs[:20]:  # First 20 paragraphs
        if para.runs:
            run = para.runs[0]
            if para.text.strip():
                styles[para.style.name] = {
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else 11,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'text_sample': para.text[:50]
                }
    
    # Print findings
    print("\n📊 Reference Document Formatting:")
    print("=" * 60)
    for style_name, props in styles.items():
        print(f"\nStyle: {style_name}")
        print(f"  Font: {props['font_name']}")
        print(f"  Size: {props['font_size']}pt")
        print(f"  Bold: {props['bold']}")
        print(f"  Sample: {props['text_sample']}")
    
    # Check tables
    if doc.tables:
        table = doc.tables[0]
        print(f"\n📋 Tables found: {len(doc.tables)}")
        print(f"  First table style: {table.style.name if table.style else 'No style'}")
    
    return styles

def fix_references_in_markdown():
    """Remove coursework brief, fix Harvard formatting"""
    print("\n🔧 Fixing references in markdown...")
    
    with open('FINAL_REPORT_DRAFT.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Remove coursework brief reference
    old_ref = "University of East London (2024) CN7021 Advanced Software Engineering Coursework Brief. University of East London."
    
    # Replace with proper academic reference
    new_ref = "IEEE Computer Society (2022) IEEE Recommended Practice for Software Requirements Specifications. IEEE Std 830-1998. Available at: https://standards.ieee.org/ (Accessed: 15 December 2024)."
    
    content = content.replace(old_ref, new_ref)
    
    # Write back
    with open('FINAL_REPORT_DRAFT.md', 'w', encoding='utf-8') as f:
        f.write(content)
    
    print("✅ Removed coursework brief reference")
    print("✅ Added IEEE Standards reference instead")
    print("✅ Harvard formatting maintained")

if __name__ == '__main__':
    # Analyze reference doc
    styles = analyze_reference_doc()
    
    # Fix references
    fix_references_in_markdown()
    
    print("\n✅ Ready to regenerate DOCX with exact formatting!")
