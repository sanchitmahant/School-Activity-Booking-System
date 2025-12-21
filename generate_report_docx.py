import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Paths
BASE_DIR = r"c:\Users\Sanchit Kaushal\OneDrive\Desktop\Advanced Software Engineering\FInal Project and Report - Copy\School_Activity_Booking_System"
TEMPLATE_PATH = os.path.join(BASE_DIR, r"Real Report Preparation\Helping and Guideline Documents\CN7021-ASWE-Project-Template-2024-25.docx")
DRAFT_PATH = os.path.join(BASE_DIR, "FINAL_REPORT_DRAFT.md")
OUTPUT_PATH = os.path.join(BASE_DIR, "Sanchit_Kaushal_Final_Report.docx")

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border via xml manipulation.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def parse_inline_formatting(paragraph, text):
    """Parses simple markdown bold and italic."""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Split by italic markers
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(subpart)

def process_table(doc, lines):
    """Creates a table from markdown lines."""
    # Filter lines
    rows = [line.strip().strip('|').split('|') for line in lines if line.strip()]
    # Remove separator line (e.g. ---|---)
    rows = [r for r in rows if not all(c.strip().replace('-', '').replace(':', '') == '' for c in r)]
    
    if not rows:
        return

    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        # Fallback: Manual borders
        pass
        
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = cell_text.strip()
                # Manual border for every cell if style failed or just to be safe
                set_cell_border(
                    cell,
                    top={"sz": 4, "val": "single", "color": "auto", "space": "0"},
                    bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"},
                    left={"sz": 4, "val": "single", "color": "auto", "space": "0"},
                    right={"sz": 4, "val": "single", "color": "auto", "space": "0"},
                    insideH={"sz": 4, "val": "single", "color": "auto", "space": "0"},
                    insideV={"sz": 4, "val": "single", "color": "auto", "space": "0"}
                )
                
                # Formatting for header
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            
    doc.add_paragraph() # Spacer after table

def fill_cover_page(doc):
    """Attempts to fill cover page placeholders."""
    replacements = {
        "Students Name & ID:": "Sanchit Kaushal [ID], Mohd Sharjeel [ID], Chichebendu Umeh [ID], Shiva Kasula [ID]",
        "Group:": "ASWE Group 1",
        "Month Year": "December 2025"
    }
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                # Naive replacement
                p.text = p.text.replace(key, f"{key} {value}")

def generate_docx():
    # Load Template
    if not os.path.exists(TEMPLATE_PATH):
        print(f"Error: Template not found at {TEMPLATE_PATH}")
        return

    doc = Document(TEMPLATE_PATH)
    fill_cover_page(doc)
    
    # Move to end of document (or finding a specific start point)
    # We will assume we append to the end.
    doc.add_page_break()

    # Read Draft
    with open(DRAFT_PATH, 'r', encoding='utf-8') as f:
        content = f.readlines()

    table_buffer = []
    in_table = False

    for line in content:
        line = line.strip()
        
        # Skip empty lines if not in table or code block
        if not line:
            if in_table:
                process_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            continue

        # Header detection
        if line.startswith('# '):
            if in_table:
                process_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            if in_table:
                process_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            if in_table:
                process_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
             if in_table:
                process_table(doc, table_buffer)
                table_buffer = []
                in_table = False
             doc.add_heading(line[5:], level=4)
             
        # Image/Placeholder detection
        elif line.startswith('[INSERT') or line.startswith('Figure '):
            if in_table:
                process_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            p = doc.add_paragraph()
            run = p.add_run(line)
            if 'INSERT' in line:
                run.font.color.rgb = RGBColor(255, 0, 0) # Red
                run.bold = True
            if 'Figure' in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.italic = True
        
        # Table detection
        elif line.startswith('|'):
            in_table = True
            table_buffer.append(line)
            
        else:
            if in_table:
                # Maybe end of table?
                if not line.startswith('|'):
                    process_table(doc, table_buffer)
                    table_buffer = []
                    in_table = False
                    p = doc.add_paragraph()
                    parse_inline_formatting(p, line)
                else:
                    table_buffer.append(line)
            else:
                p = doc.add_paragraph()
                parse_inline_formatting(p, line)

    # Flush table if any
    if in_table and table_buffer:
        process_table(doc, table_buffer)

    doc.save(OUTPUT_PATH)
    print(f"Generated report at: {OUTPUT_PATH}")

if __name__ == "__main__":
    generate_docx()
