import csv
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_FILE = 'documentation/Project_Plan_Updated_2025.csv'
OUTPUT_DIR = 'documentation/individual contributions'

# Mapping tasks to relevant files for code explanation
TASK_FILE_MAPPING = {
    "Phase 1: Inception & Scoping": ["app.py", "requirements.txt"],
    "Requirements Engineering": ["documentation/Announcement.docx", "documentation/CN7021 ASWE 2025-26 Coursework Brief.docx"],
    "Phase 2: Architecture & Database Design": ["app.py", "config.py"],
    "Database Schema Design (3NF)": ["app.py", "populate_db.py"],
    "Security Architecture": ["app.py", "config.py"],
    "UI/UX Design System": ["static/style.css", "templates/base.html"],
    "Phase 3: Core Authentication System": ["app.py", "templates/login.html", "templates/register.html"],
    "Role-Based Access Control (RBAC)": ["app.py"],
    "Parent Portal Development": ["app.py", "templates/dashboard.html"],
    "Phase 4: Advanced Booking Engine": ["app.py", "templates/activity_enrollments.html"],
    "Database Transaction Management": ["app.py"],
    "Waitlist System (Automated)": ["app.py"],
    "Payment Flow Implementation": ["app.py", "templates/payment.html"],
    "Phase 5: Advanced Features Integration": ["app.py", "requirements.txt"],
    "Email Notification System": ["app.py", "config.py"],
    "Calendar Integration (.ics Files)": ["app.py"],
    "PDF Invoice Generation": ["enhanced_invoice.py", "app.py"],
    "Admin Portal Development": ["app.py", "templates/admin/dashboard.html"],
    "Tutor Portal Development": ["app.py", "templates/tutor/dashboard.html"],
    "≡ƒåò Tutor Registration & Approval System": ["app.py", "templates/tutor/register.html", "templates/admin/pending_tutors.html"],
    "Phase 6: Deployment Preparation": ["requirements.txt", "config.py", "README.md"],
    "Testing & Quality Assurance": ["tests/"],
    "Documentation & Academic Compliance": ["documentation/"]
}

# Viva Questions per task
VIVA_QUESTIONS = {
    "Phase 1: Inception & Scoping": [
        "Why did you choose Flask over Django for this project?",
        "Explain the Application Factory Pattern and why it's useful.",
        "How does the MVC architecture apply to your Flask application?"
    ],
    "Requirements Engineering": [
        "How did you identify the three user personas?",
        "What was the most critical requirement you discovered during analysis?",
        "How did you prioritize features for the MVP?"
    ],
    "Phase 2: Architecture & Database Design": [
        "Explain the benefits of the MVC pattern in your architecture.",
        "How does your architecture support scalability?",
        "Why did you choose SQLAlchemy as your ORM?"
    ],
    "Database Schema Design (3NF)": [
        "What is Third Normal Form (3NF) and how did you apply it?",
        "Explain the relationship between the Parent and Child tables.",
        "How do you handle data integrity in your schema?"
    ],
    "Security Architecture": [
        "How does CSRF protection work in your application?",
        "Why did you use Scrypt for password hashing?",
        "How do you prevent SQL injection attacks?"
    ],
    "UI/UX Design System": [
        "How did you ensure accessibility (WCAG) in your design?",
        "Explain the responsive grid system you implemented.",
        "Why did you choose this specific color palette?"
    ],
    "Phase 3: Core Authentication System": [
        "Walk me through the user registration flow.",
        "How do you securely manage user sessions?",
        "What happens if a user tries to access a protected route without logging in?"
    ],
    "Role-Based Access Control (RBAC)": [
        "How do your custom decorators enforce RBAC?",
        "Why is RBAC better than simple boolean flags for permissions?",
        "Can a user have multiple roles in your system?"
    ],
    "Parent Portal Development": [
        "How do you calculate real-time activity availability?",
        "Explain the AJAX implementation for adding a child.",
        "How does the dashboard handle database queries efficiently?"
    ],
    "Phase 4: Advanced Booking Engine": [
        "How do you detect temporal conflicts during booking?",
        "Explain the validation pipeline for a new booking.",
        "How do you handle race conditions when capacity is low?"
    ],
    "Database Transaction Management": [
        "What are ACID properties and how do they apply here?",
        "Show me where you use database transactions in the code.",
        "What happens if an error occurs during a multi-step database operation?"
    ],
    "Waitlist System (Automated)": [
        "Explain the algorithm for automatic waitlist promotion.",
        "How do you ensure fairness in the waitlist queue?",
        "What triggers the promotion of a student from the waitlist?"
    ],
    "Payment Flow Implementation": [
        "Explain the state machine pattern used in your payment flow.",
        "How would you integrate a real payment gateway like Stripe?",
        "Why do you separate the payment confirmation from the booking creation?"
    ],
    "Phase 5: Advanced Features Integration": [
        "What challenges did you face integrating external libraries?",
        "How do you manage dependencies for these advanced features?",
        "Which feature added the most value to the project?"
    ],
    "Email Notification System": [
        "How do you send emails asynchronously to avoid blocking the UI?",
        "Explain the configuration required for Gmail SMTP.",
        "How do you handle email sending failures?"
    ],
    "Calendar Integration (.ics Files)": [
        "What is the iCalendar standard (RFC 5545)?",
        "How do you generate the unique content for the .ics file?",
        "How does the user interact with the generated calendar file?"
    ],
    "PDF Invoice Generation": [
        "Why did you use ReportLab instead of HTML-to-PDF conversion?",
        "Explain how you draw tables and styles programmatically.",
        "How do you ensure the PDF is downloadable rather than just displayed?"
    ],
    "Admin Portal Development": [
        "How did you implement the CRUD operations for activities?",
        "Explain the cascade delete behavior when removing a tutor.",
        "How do you calculate the real-time statistics on the dashboard?"
    ],
    "Tutor Portal Development": [
        "How does the batch attendance processing work efficiently?",
        "How do you ensure tutors only see their assigned activities?",
        "Explain the database query for fetching tutor schedules."
    ],
    "≡ƒåò Tutor Registration & Approval System": [
        "Walk me through the tutor approval workflow.",
        "How do you secure the admin approval endpoints?",
        "What changes did you make to the database to support this feature?"
    ],
    "Phase 6: Deployment Preparation": [
        "What is the purpose of the Procfile?",
        "How do you manage environment variables in production?",
        "Explain the difference between the development and production configurations."
    ],
    "Testing & Quality Assurance": [
        "What testing strategies did you employ?",
        "How did you perform cross-browser testing?",
        "What was the most critical bug you found and fixed?"
    ],
    "Documentation & Academic Compliance": [
        "How did you calculate the COCOMO metrics?",
        "What software metrics did you use to assess code quality?",
        "Explain the structure of your technical report."
    ]
}

def create_docx(assignee, tasks):
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Individual Contribution Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading(f"Contributor: {assignee}", level=1)
    doc.add_paragraph(f"Role: Full Stack Developer & Team Member")
    
    # Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"{assignee} was responsible for {len(tasks)} key components of the School Activity Booking System. Their contributions spanned across {', '.join(set([t['Name'].split(':')[0] for t in tasks]))}, demonstrating proficiency in full-stack development, database management, and system architecture.")
    
    # Tasks
    doc.add_heading("Detailed Contributions", level=1)
    
    for i, task in enumerate(tasks, 1):
        task_name = task['Name']
        doc.add_heading(f"{i}. {task_name}", level=2)
        
        # Simple Explanation
        doc.add_heading("Simple Explanation (Non-Technical)", level=3)
        simple_text = f"In this task, I {task['Description'].lower()}. This feature allows users to easily interact with the system by providing a seamless experience. It solves the problem of manual management by automating the process."
        doc.add_paragraph(simple_text, style='Intense Quote')
        
        # Technical Explanation
        doc.add_heading("Technical Explanation (Detailed)", level=3)
        tech_text = task.get('Detailed Presentation Script (Read this for Viva/Questions)', 'No detailed description available.')
        doc.add_paragraph(tech_text)
        
        # Code Logic
        doc.add_heading("Code Logic & Implementation Details", level=3)
        files = TASK_FILE_MAPPING.get(task_name, [])
        if files:
            doc.add_paragraph(f"Key Files Involved: {', '.join(files)}")
            doc.add_paragraph("Implementation Strategy:")
            doc.add_paragraph("1. Defined data models in app.py using SQLAlchemy.")
            doc.add_paragraph("2. Created route handlers (controllers) to process requests.")
            doc.add_paragraph("3. Implemented frontend templates using Jinja2 and Bootstrap.")
            doc.add_paragraph("4. Integrated validation logic to ensure data integrity.")
        else:
            doc.add_paragraph("This task involved high-level planning and documentation, focusing on architectural decisions rather than direct code implementation.")
            
        # Viva Questions
        doc.add_heading("Potential Viva Questions & Answers", level=3)
        questions = VIVA_QUESTIONS.get(task_name, ["Explain this feature in detail.", "What were the challenges faced?", "How would you improve this?"])
        for q in questions:
            p = doc.add_paragraph()
            runner = p.add_run(f"Q: {q}")
            runner.bold = True
            doc.add_paragraph("A: [Prepare your answer based on the technical explanation above]")
            
    return doc

def create_markdown(assignee, tasks):
    md = []
    md.append(f"# Individual Contribution Report: {assignee}")
    md.append(f"**Role:** Full Stack Developer & Team Member\n")
    
    md.append("## Executive Summary")
    md.append(f"{assignee} was responsible for {len(tasks)} key components of the School Activity Booking System. Their contributions spanned across {', '.join(set([t['Name'].split(':')[0] for t in tasks]))}, demonstrating proficiency in full-stack development, database management, and system architecture.\n")
    
    md.append("## Detailed Contributions\n")
    
    for i, task in enumerate(tasks, 1):
        task_name = task['Name']
        md.append(f"### {i}. {task_name}")
        
        md.append("#### Simple Explanation (Non-Technical)")
        md.append(f"> In this task, I {task['Description'].lower()}. This feature allows users to easily interact with the system by providing a seamless experience. It solves the problem of manual management by automating the process.\n")
        
        md.append("#### Technical Explanation (Detailed)")
        tech_text = task.get('Detailed Presentation Script (Read this for Viva/Questions)', 'No detailed description available.')
        md.append(f"{tech_text}\n")
        
        md.append("#### Code Logic & Implementation Details")
        files = TASK_FILE_MAPPING.get(task_name, [])
        if files:
            md.append(f"**Key Files Involved:** `{', '.join(files)}`\n")
            md.append("**Implementation Strategy:**")
            md.append("1. Defined data models in `app.py` using SQLAlchemy.")
            md.append("2. Created route handlers (controllers) to process requests.")
            md.append("3. Implemented frontend templates using Jinja2 and Bootstrap.")
            md.append("4. Integrated validation logic to ensure data integrity.\n")
        else:
            md.append("This task involved high-level planning and documentation, focusing on architectural decisions rather than direct code implementation.\n")
            
        md.append("#### Potential Viva Questions")
        questions = VIVA_QUESTIONS.get(task_name, ["Explain this feature in detail.", "What were the challenges faced?", "How would you improve this?"])
        for q in questions:
            md.append(f"- **Q:** {q}")
            md.append("  - **A:** *[Prepare your answer based on the technical explanation above]*")
        md.append("\n---\n")
        
    return "\n".join(md)

def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        
    # Read CSV
    tasks_by_assignee = {}
    with open(INPUT_FILE, 'r', encoding='utf-16') as f:
        reader = csv.DictReader(f)
        for row in reader:
            assignee = row['Assignee'].strip()
            if assignee not in tasks_by_assignee:
                tasks_by_assignee[assignee] = []
            tasks_by_assignee[assignee].append(row)
            
    # Generate Docs
    for assignee, tasks in tasks_by_assignee.items():
        safe_name = assignee.replace(" ", "_")
        
        # Generate DOCX
        doc = create_docx(assignee, tasks)
        docx_path = os.path.join(OUTPUT_DIR, f"{safe_name}_Contribution.docx")
        doc.save(docx_path)
        print(f"Generated DOCX: {docx_path}")
        
        # Generate Markdown
        md_content = create_markdown(assignee, tasks)
        md_path = os.path.join(OUTPUT_DIR, f"{safe_name}_Contribution.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        print(f"Generated MD: {md_path}")

if __name__ == "__main__":
    main()
